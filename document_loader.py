"""
Document loader module for PDF and Word document processing.
Handles loading and chunking of documents from folders.
Supports multiple files per module with proper metadata tracking.
Supported formats: PDF, DOC, DOCX
"""

import os
from pathlib import Path
from typing import List, Dict, Any, Generator

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document

from config import ChunkingConfig

# Supported file extensions
SUPPORTED_EXTENSIONS = {'.pdf', '.doc', '.docx'}


class PDFLoadError(Exception):
    """Custom exception for document loading errors."""
    pass


class DocumentLoadError(Exception):
    """Custom exception for document loading errors."""
    pass


class FolderDocumentProcessor:
    """
    Handles document loading and text chunking from a folder.
    Processes all supported files (PDF, DOC, DOCX) in a module's document folder.
    
    Attributes:
        folder_path: Path to the folder containing document files
        module_id: Identifier for the module
        chunk_size: Size of each text chunk
        chunk_overlap: Overlap between consecutive chunks
    """
    
    def __init__(
        self,
        folder_path: str,
        module_id: str,
        chunk_size: int = ChunkingConfig.CHUNK_SIZE,
        chunk_overlap: int = ChunkingConfig.CHUNK_OVERLAP
    ):
        """
        Initialize the folder document processor.
        
        Args:
            folder_path: Path to the folder containing PDF files
            module_id: Identifier for the module
            chunk_size: Maximum size of each chunk
            chunk_overlap: Number of characters to overlap between chunks
        """
        self.folder_path = Path(folder_path)
        self.module_id = module_id
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self._text_splitter = self._create_text_splitter()
    
    def _create_text_splitter(self) -> RecursiveCharacterTextSplitter:
        """Create and configure the text splitter."""
        return RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ".", " ", ""]
        )
    
    def ensure_folder_exists(self):
        """Create the folder if it doesn't exist."""
        self.folder_path.mkdir(parents=True, exist_ok=True)
    
    def get_pdf_files(self) -> List[Path]:
        """Get all supported document files in the folder (PDF, DOC, DOCX)."""
        if not self.folder_path.exists():
            return []
        files = []
        for ext in SUPPORTED_EXTENSIONS:
            files.extend(self.folder_path.glob(f"*{ext}"))
        return sorted(files, key=lambda x: x.name.lower())
    
    def get_file_list(self) -> List[Dict[str, Any]]:
        """Get list of PDF files with metadata."""
        files = []
        for pdf_path in self.get_pdf_files():
            stat = pdf_path.stat()
            files.append({
                "name": pdf_path.name,
                "path": str(pdf_path),
                "size": stat.st_size,
                "modified": stat.st_mtime
            })
        return files
    
    def _load_word_document(self, doc_path: Path) -> List[Document]:
        """
        Load a Word document (.doc or .docx).
        
        Args:
            doc_path: Path to the Word document
            
        Returns:
            List[Document]: List of document content
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise DocumentLoadError(
                "python-docx n'est pas installé. Exécutez: pip install python-docx"
            )
        
        try:
            doc = DocxDocument(str(doc_path))
            
            # Extract text from paragraphs
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            content = "\n\n".join(full_text)
            
            # Create single document (Word doesn't have pages like PDF)
            return [Document(
                page_content=content,
                metadata={
                    "file_name": doc_path.name,
                    "module": self.module_id,
                    "page": 1,
                    "source": str(doc_path)
                }
            )]
            
        except Exception as e:
            raise DocumentLoadError(
                f"Erreur lors du chargement de '{doc_path.name}': {str(e)}"
            )
    
    def load_single_pdf(self, pdf_path: Path) -> List[Document]:
        """
        Load a single document file (PDF, DOC, or DOCX).
        
        Args:
            pdf_path: Path to the document file
            
        Returns:
            List[Document]: List of document pages/content
        """
        file_ext = pdf_path.suffix.lower()
        
        try:
            if file_ext == '.pdf':
                loader = PyPDFLoader(str(pdf_path))
                documents = loader.load()
            elif file_ext in ['.doc', '.docx']:
                documents = self._load_word_document(pdf_path)
            else:
                raise DocumentLoadError(
                    f"Format non supporté: {file_ext}. Formats supportés: PDF, DOC, DOCX"
                )
            
            # Add file name to metadata
            file_name = pdf_path.name
            for doc in documents:
                doc.metadata["file_name"] = file_name
                doc.metadata["module"] = self.module_id
            
            return documents
            
        except (PDFLoadError, DocumentLoadError):
            raise
        except Exception as e:
            raise DocumentLoadError(
                f"Erreur lors du chargement de '{pdf_path.name}': {str(e)}"
            )
    
    def split_documents(self, documents: List[Document], file_name: str) -> List[Document]:
        """
        Split documents into smaller chunks with proper metadata.
        
        Args:
            documents: List of documents to split
            file_name: Name of the source file
            
        Returns:
            List[Document]: List of chunked documents
        """
        chunks = self._text_splitter.split_documents(documents)
        
        for i, chunk in enumerate(chunks):
            chunk.metadata["chunk_id"] = i
            chunk.metadata["file_name"] = file_name
            chunk.metadata["module"] = self.module_id
        
        return chunks
    
    def process_single_file(self, pdf_path: Path) -> List[Document]:
        """
        Load and process a single PDF file.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            List[Document]: Chunked documents from the file
        """
        documents = self.load_single_pdf(pdf_path)
        return self.split_documents(documents, pdf_path.name)
    
    def process_all_files(self) -> Generator[Dict[str, Any], None, None]:
        """
        Process all PDF files in the folder.
        
        Yields:
            Dict with file info and processed chunks
        """
        pdf_files = self.get_pdf_files()
        
        for pdf_path in pdf_files:
            try:
                chunks = self.process_single_file(pdf_path)
                yield {
                    "file_name": pdf_path.name,
                    "file_path": str(pdf_path),
                    "chunks": chunks,
                    "chunk_count": len(chunks),
                    "success": True,
                    "error": None
                }
            except PDFLoadError as e:
                yield {
                    "file_name": pdf_path.name,
                    "file_path": str(pdf_path),
                    "chunks": [],
                    "chunk_count": 0,
                    "success": False,
                    "error": str(e)
                }
    
    def load_all_documents(self) -> List[Document]:
        """
        Load and chunk all PDF files in the folder.
        
        Returns:
            List[Document]: All chunked documents from all files
        """
        all_chunks = []
        
        for result in self.process_all_files():
            if result["success"]:
                all_chunks.extend(result["chunks"])
        
        return all_chunks


# Keep backward compatibility with single file processor
class DocumentProcessor(FolderDocumentProcessor):
    """
    Single file document processor for backward compatibility.
    """
    
    def __init__(
        self,
        pdf_path: str,
        chunk_size: int = ChunkingConfig.CHUNK_SIZE,
        chunk_overlap: int = ChunkingConfig.CHUNK_OVERLAP
    ):
        self.pdf_path = Path(pdf_path)
        self.module_id = "default"
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self._text_splitter = self._create_text_splitter()
    
    def validate_pdf_exists(self) -> bool:
        """Check if the PDF file exists."""
        return self.pdf_path.exists()
    
    def load_and_split(self) -> List[Document]:
        """Load PDF and split into chunks."""
        if not self.validate_pdf_exists():
            raise PDFLoadError(
                f"Le fichier PDF '{self.pdf_path}' n'a pas été trouvé."
            )
        return self.process_single_file(self.pdf_path)


def create_folder_processor(module_config: Dict[str, Any]) -> FolderDocumentProcessor:
    """
    Factory function to create a FolderDocumentProcessor for a specific module.
    
    Args:
        module_config: Module configuration dictionary
        
    Returns:
        FolderDocumentProcessor: Configured processor instance
    """
    return FolderDocumentProcessor(
        folder_path=module_config["documents_folder"],
        module_id=module_config["id"]
    )


# Backward compatibility
def create_document_processor(module_config: Dict[str, Any]) -> FolderDocumentProcessor:
    """Alias for create_folder_processor for backward compatibility."""
    return create_folder_processor(module_config)
